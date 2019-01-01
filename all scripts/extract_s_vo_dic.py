from modules.enhanced_svo_extraction.subject_verb_object_extract import findSVOs, nlp
from nltk.tokenize import word_tokenize,sent_tokenize
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer
from spacy.lang.en import English
import nltk
import re
import docx


# preparing data set
# docx_file = docx.Document('test_files/project_abstracts.docx')
# print(len(docx_file.paragraphs))
# abstracts = list()
# for i in range(int(len(docx_file.paragraphs)/3)):
#     abstracts.append((docx_file.paragraphs[i*3].text,docx_file.paragraphs[i*3+1].text,docx_file.paragraphs[i*3+2].text))


# function to print formatted dictionary
def print_formatted_dict(dic,spacing=40):
    print("DATA".ljust(spacing),"FREQUENCY".ljust(spacing))
    for key in dic:
        print(str(key).ljust(spacing),str(dic[key]).ljust(spacing))


stop_words = set(stopwords.words("english"))
lemmatizer = WordNetLemmatizer()
noun_chunk_gram = r"""Chunk: {<NN.*>*}"""
noun_parser = nltk.RegexpParser(noun_chunk_gram)

# looping through each data_set

# text = open("test_files/4.txt","r").read()
text = "Abstract thinking is a level of thinking about things that is removed from the facts of the “here and now”, and from specific examples of the things or concepts being thought about. Abstract thinkers are able to reflect on events and ideas, and on attributes and relationships separate from the objects that have those attributes or share those relationships. Thus, for example, a concrete thinker can think about this particular dog; a more abstract thinker can think about dogs in general. A concrete thinker can think about this dog on this rug; a more abstract thinker can think about spatial relations, like “on”. A concrete thinker can see that this ball is big; a more abstract thinker can think about size in general. A concrete thinker can count three cookies; a more abstract thinker can think about numbers. A concrete thinker can recognize that John likes Betty; a more abstract thinker can reflect on emotions, like affection."
abstracts = [text]

for i,abstract in enumerate(abstracts):
    # title, abstract, keywords = data_tup
    # print(i+1,".",title)

    noun_dic = dict()
    vo_dic = dict()

    sents = sent_tokenize(abstract.lower())

    for sent in sents:
        # print(sent)
        tagged = nltk.pos_tag(word_tokenize(sent))
        # print(tagged)

        # finding svo(s)
        tok = nlp(sent)
        svos = findSVOs(tok)
        # print(svos)

        # filtering and lemmatizing verb-object
        for subj,ver,obj in svos:
            ver = lemmatizer.lemmatize(ver,"v")
            tag_words = nltk.pos_tag(word_tokenize(obj))
            for word, tag in tag_words:
                if word in stop_words:
                    obj = re.sub("\\b" + word + "\\b", "", obj)
                elif not re.match(r"[A-Z]+", tag):
                    obj = obj.replace(word, "")
            obj = obj.strip()
            obj = obj.replace("  ", " ")
            if not obj.isspace() and len(obj) > 0:
                vo_dic[(ver,obj)] = vo_dic.get((ver,obj),0)+1
        # print(vo_dic)
        # nouns = dict()

        # finding nouns
        chunked_n = noun_parser.parse(tagged)
        # chunked.draw()
        # printing subjects
        for subtree in chunked_n.subtrees(filter=lambda t :t.label() =='Chunk'):
            subs = [ word for word,tag in subtree.leaves()]
            term = lemmatizer.lemmatize(" ".join(subs))
            noun_dic[term] = noun_dic.get(term,0)+1

        # print("\n\n\n")

    print("\nVerb-Object Pairs....")
    print_formatted_dict(vo_dic,40)
    # print(vo_dic)
    print("\nNouns....")
    print_formatted_dict(noun_dic,30)
    # print(noun_dic)
    print("\n\n\n")
